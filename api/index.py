import os
import json
import datetime
from flask import Flask, request, jsonify, render_template, send_file
from openai import OpenAI
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from dotenv import load_dotenv

app = Flask(__name__, template_folder="../templates", static_folder="../static")

load_dotenv()

client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
HISTORY_FILE = "history.json"

generated_cache = {"paper": ""}


def save_history(entry):
    try:
        with open(HISTORY_FILE, "r") as f:
            history = json.load(f)
    except:
        history = []

    history.insert(0, entry)

    with open(HISTORY_FILE, "w") as f:
        json.dump(history[:50], f)


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/api/generate-paper", methods=["POST"])
def generate_paper():
    data = request.json

    subjects = data.get("subjects", [])
    class_level = data.get("class")
    board = data.get("board")
    marks = data.get("marks")
    difficulty = data.get("difficulty")
    instructions = data.get("instructions", "")

    prompt = f"""
You are a senior {board} examination paper setter.

Generate a professional exam paper.

Class: {class_level}
Board: {board}
Subjects: {", ".join(subjects)}
Marks: {marks}
Difficulty: {difficulty}

Custom Instructions:
{instructions}

Include sections:
Section A: MCQ
Section B: Short Answer
Section C: Long Answer
Section D: Case Study

Ensure exact marks total.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Expert question paper generator."},
            {"role": "user", "content": prompt}
        ]
    )

    paper = response.choices[0].message.content

    generated_cache["paper"] = paper

    entry = {
        "date": str(datetime.datetime.now()),
        "class": class_level,
        "board": board,
        "subjects": subjects,
        "marks": marks,
        "difficulty": difficulty,
        "paper": paper
    }

    save_history(entry)

    return jsonify({"paper": paper})


@app.route("/api/history")
def history():
    try:
        with open(HISTORY_FILE) as f:
            return jsonify(json.load(f))
    except:
        return jsonify([])


@app.route("/api/download-pdf")
def download_pdf():
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    text = c.beginText(40, 750)

    for line in generated_cache["paper"].split("\n"):
        text.textLine(line)

    c.drawText(text)
    c.save()

    buffer.seek(0)

    return send_file(buffer, as_attachment=True,
                     download_name="question_paper.pdf",
                     mimetype="application/pdf")


@app.route("/api/download-word")
def download_word():
    document = Document()
    document.add_heading("Question Paper", 0)
    document.add_paragraph(generated_cache["paper"])

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True,
                     download_name="question_paper.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


app = app